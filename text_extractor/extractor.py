"""
Text Extractor

Extracts text from PDF, DOCX, DOC, TXT files.
Output is in markdown format.
Embedded images are saved to folder with paths in markdown.
"""

import os
import uuid
import logging
from typing import List, Optional, Tuple
from pathlib import Path
from io import BytesIO

import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

from .schemas import (
    FileType,
    ContentType,
    ExtractedBlock,
    DocumentMetadata,
    ExtractionRequest,
    ExtractionResult
)

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extracts text from multiple file formats.
    Output is markdown formatted.
    Images are saved to specified folder.

    Supported formats: PDF, DOCX, DOC, TXT
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MIN_IMAGE_WIDTH = 50
    MIN_IMAGE_HEIGHT = 50

    def __init__(self):
        self.sequence_counter = 0
        self.image_counter = 0
        self.image_paths: List[str] = []

    def extract(self, request: ExtractionRequest) -> ExtractionResult:
        """
        Extract text from document.

        Args:
            request: ExtractionRequest with file path and options

        Returns:
            ExtractionResult with markdown text and image paths
        """
        file_path = Path(request.file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        # Generate document ID
        document_id = request.document_id or str(uuid.uuid4())
        self.sequence_counter = 0
        self.image_counter = 0
        self.image_paths = []

        # Setup image output directory
        if request.include_images:
            if request.image_output_dir:
                image_dir = Path(request.image_output_dir)
            else:
                image_dir = file_path.parent / f"{document_id}_images"
            image_dir.mkdir(parents=True, exist_ok=True)
        else:
            image_dir = None

        # Get file info
        file_size = file_path.stat().st_size
        file_type = self._get_file_type(extension)

        # Extract based on file type
        if extension == '.pdf':
            blocks, total_pages = self._extract_pdf(
                file_path, document_id, request.include_tables,
                request.include_images, image_dir
            )
        elif extension == '.docx':
            blocks, total_pages = self._extract_docx(
                file_path, document_id, request.include_tables,
                request.include_images, image_dir
            )
        elif extension == '.doc':
            blocks, total_pages = self._extract_doc(
                file_path, document_id, request.include_tables
            )
        elif extension == '.txt':
            blocks, total_pages = self._extract_txt(file_path, document_id)
        else:
            raise ValueError(f"Unsupported extension: {extension}")

        # Build markdown output
        markdown_text = self._build_markdown(blocks)

        # Calculate stats
        word_count = sum(len(b.text.split()) for b in blocks if b.content_type != ContentType.IMAGE)
        char_count = sum(len(b.text) for b in blocks if b.content_type != ContentType.IMAGE)

        # Create metadata
        metadata = DocumentMetadata(
            document_id=document_id,
            filename=file_path.name,
            file_type=file_type,
            file_size_bytes=file_size,
            total_pages=total_pages,
            total_blocks=len(blocks),
            word_count=word_count,
            char_count=char_count,
            status="completed"
        )

        return ExtractionResult(
            metadata=metadata,
            blocks=blocks,
            markdown_text=markdown_text,
            image_paths=self.image_paths
        )

    def _get_file_type(self, extension: str) -> FileType:
        """Map extension to FileType enum."""
        mapping = {
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.doc': FileType.DOC,
            '.txt': FileType.TXT
        }
        return mapping.get(extension, FileType.TXT)

    def _create_block(
        self,
        document_id: str,
        content_type: ContentType,
        text: str,
        page: Optional[int] = None,
        level: Optional[int] = None,
        metadata: dict = None
    ) -> ExtractedBlock:
        """Create an ExtractedBlock with auto-incremented sequence."""
        block = ExtractedBlock(
            block_id=f"{document_id}_block_{self.sequence_counter:05d}",
            document_id=document_id,
            sequence=self.sequence_counter,
            content_type=content_type,
            text=text.strip() if content_type != ContentType.IMAGE else text,
            page=page,
            level=level,
            metadata=metadata or {}
        )
        self.sequence_counter += 1
        return block

    def _save_image(
        self,
        image_bytes: bytes,
        image_dir: Path,
        document_id: str,
        page: Optional[int],
        ext: str = "png"
    ) -> str:
        """Save image to folder and return path."""
        self.image_counter += 1

        if page:
            filename = f"{document_id}_page{page}_img{self.image_counter:03d}.{ext}"
        else:
            filename = f"{document_id}_img{self.image_counter:03d}.{ext}"

        image_path = image_dir / filename

        with open(image_path, 'wb') as f:
            f.write(image_bytes)

        self.image_paths.append(str(image_path))
        return str(image_path)

    # =====================
    # PDF Extraction
    # =====================

    def _extract_pdf(
        self,
        file_path: Path,
        document_id: str,
        include_tables: bool,
        include_images: bool,
        image_dir: Optional[Path]
    ) -> Tuple[List[ExtractedBlock], int]:
        """Extract text and images from PDF using PyMuPDF."""
        blocks = []
        doc = fitz.open(str(file_path))

        try:
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                page_blocks = self._extract_pdf_page(
                    page, doc, document_id, page_num + 1,
                    include_tables, include_images, image_dir
                )
                blocks.extend(page_blocks)

            return blocks, total_pages

        finally:
            doc.close()

    def _extract_pdf_page(
        self,
        page: fitz.Page,
        doc: fitz.Document,
        document_id: str,
        page_num: int,
        include_tables: bool,
        include_images: bool,
        image_dir: Optional[Path]
    ) -> List[ExtractedBlock]:
        """Extract blocks from a single PDF page maintaining order."""
        elements = []  # (y_position, x_position, block)

        # Get text blocks with structure
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        for block in text_dict.get("blocks", []):
            bbox = block.get("bbox", [0, 0, 0, 0])
            y_pos = bbox[1]
            x_pos = bbox[0]

            if block.get("type") == 0:  # Text block
                block_text, is_heading, heading_level = self._process_pdf_block(block)

                if block_text.strip():
                    content_type = ContentType.HEADING if is_heading else ContentType.PARAGRAPH

                    extracted = self._create_block(
                        document_id=document_id,
                        content_type=content_type,
                        text=block_text,
                        page=page_num,
                        level=heading_level if is_heading else None,
                        metadata={"bbox": bbox}
                    )
                    elements.append((y_pos, x_pos, extracted))

        # Extract images
        if include_images and image_dir:
            image_list = page.get_images(full=True)

            for img_info in image_list:
                xref = img_info[0]

                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue

                    image_bytes = base_image["image"]
                    width = base_image.get("width", 0)
                    height = base_image.get("height", 0)
                    ext = base_image.get("ext", "png")

                    # Filter small images
                    if width < self.MIN_IMAGE_WIDTH or height < self.MIN_IMAGE_HEIGHT:
                        continue

                    # Get image position
                    img_rects = page.get_image_rects(xref)
                    if img_rects:
                        rect = img_rects[0]
                        y_pos = rect.y0
                        x_pos = rect.x0
                        bbox = [rect.x0, rect.y0, rect.x1, rect.y1]
                    else:
                        y_pos = 0
                        x_pos = 0
                        bbox = [0, 0, width, height]

                    # Save image
                    image_path = self._save_image(
                        image_bytes, image_dir, document_id, page_num, ext
                    )

                    # Create markdown image reference
                    md_image = f"![image]({image_path})"

                    extracted = self._create_block(
                        document_id=document_id,
                        content_type=ContentType.IMAGE,
                        text=md_image,
                        page=page_num,
                        metadata={
                            "bbox": bbox,
                            "width": width,
                            "height": height,
                            "image_path": image_path
                        }
                    )
                    elements.append((y_pos, x_pos, extracted))

                except Exception as e:
                    logger.warning(f"Failed to extract image on page {page_num}: {e}")

        # Sort by position (top to bottom, left to right)
        elements.sort(key=lambda x: (x[0], x[1]))

        return [elem[2] for elem in elements]

    def _process_pdf_block(self, block: dict) -> Tuple[str, bool, Optional[int]]:
        """Process PDF text block and detect if it's a heading."""
        lines = []
        max_font_size = 0
        is_bold = False

        for line in block.get("lines", []):
            line_text = ""
            for span in line.get("spans", []):
                line_text += span.get("text", "")
                font_size = span.get("size", 12)
                if font_size > max_font_size:
                    max_font_size = font_size
                if "bold" in span.get("font", "").lower():
                    is_bold = True

            if line_text.strip():
                lines.append(line_text)

        text = "\n".join(lines)

        # Heading detection
        is_heading = False
        heading_level = None

        if max_font_size >= 16 or is_bold:
            if len(text) < 200 and not text.endswith('.'):
                is_heading = True
                if max_font_size >= 24:
                    heading_level = 1
                elif max_font_size >= 18:
                    heading_level = 2
                else:
                    heading_level = 3

        return text, is_heading, heading_level

    # =====================
    # DOCX Extraction
    # =====================

    def _extract_docx(
        self,
        file_path: Path,
        document_id: str,
        include_tables: bool,
        include_images: bool,
        image_dir: Optional[Path]
    ) -> Tuple[List[ExtractedBlock], int]:
        """Extract text and images from DOCX."""
        blocks = []
        doc = DocxDocument(str(file_path))

        # Extract images from docx
        if include_images and image_dir:
            self._extract_docx_images(doc, document_id, image_dir)

        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = self._find_paragraph_by_element(doc, element)
                if para:
                    # Check for inline images
                    inline_images = self._get_inline_images(para, doc, document_id, image_dir, include_images)

                    block = self._process_docx_paragraph(para, document_id)
                    if block and block.text.strip():
                        blocks.append(block)

                    # Add inline images after paragraph
                    blocks.extend(inline_images)

            elif element.tag.endswith('tbl') and include_tables:
                table = self._find_table_by_element(doc, element)
                if table:
                    block = self._process_docx_table(table, document_id)
                    if block:
                        blocks.append(block)

        return blocks, None

    def _extract_docx_images(
        self,
        doc: DocxDocument,
        document_id: str,
        image_dir: Path
    ):
        """Extract embedded images from DOCX."""
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    ext = rel.target_part.content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'
                    self._save_image(image_data, image_dir, document_id, None, ext)
                except Exception as e:
                    logger.warning(f"Failed to extract DOCX image: {e}")

    def _get_inline_images(
        self,
        para,
        doc: DocxDocument,
        document_id: str,
        image_dir: Optional[Path],
        include_images: bool
    ) -> List[ExtractedBlock]:
        """Get inline images from paragraph."""
        blocks = []

        if not include_images or not image_dir:
            return blocks

        # Check for drawing elements (images)
        drawings = para._element.findall('.//' + qn('w:drawing'))

        for drawing in drawings:
            blips = drawing.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed:
                    try:
                        rel = doc.part.rels[embed]
                        if "image" in rel.reltype:
                            image_data = rel.target_part.blob
                            ext = rel.target_part.content_type.split('/')[-1]
                            if ext == 'jpeg':
                                ext = 'jpg'

                            image_path = self._save_image(
                                image_data, image_dir, document_id, None, ext
                            )

                            md_image = f"![image]({image_path})"
                            blocks.append(self._create_block(
                                document_id=document_id,
                                content_type=ContentType.IMAGE,
                                text=md_image,
                                metadata={"image_path": image_path}
                            ))
                    except Exception as e:
                        logger.warning(f"Failed to extract inline image: {e}")

        return blocks

    def _find_paragraph_by_element(self, doc: DocxDocument, element) -> Optional[any]:
        """Find paragraph object by XML element."""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_by_element(self, doc: DocxDocument, element) -> Optional[any]:
        """Find table object by XML element."""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _process_docx_paragraph(self, para, document_id: str) -> Optional[ExtractedBlock]:
        """Process a DOCX paragraph."""
        text = para.text.strip()
        if not text:
            return None

        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "Title" in style_name
        heading_level = None

        if is_heading:
            for i in range(1, 7):
                if f"Heading {i}" in style_name:
                    heading_level = i
                    break
            if "Title" in style_name:
                heading_level = 1

        is_list = self._is_docx_list_item(para)

        if is_heading:
            content_type = ContentType.HEADING
        elif is_list:
            content_type = ContentType.LIST_ITEM
        else:
            content_type = ContentType.PARAGRAPH

        return self._create_block(
            document_id=document_id,
            content_type=content_type,
            text=text,
            level=heading_level,
            metadata={"style": style_name}
        )

    def _is_docx_list_item(self, para) -> bool:
        """Check if paragraph is a list item."""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            return numPr is not None
        return False

    def _process_docx_table(self, table: DocxTable, document_id: str) -> Optional[ExtractedBlock]:
        """Process a DOCX table and convert to markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append(cells)

        if not rows:
            return None

        md_table = self._table_to_markdown(rows)

        return self._create_block(
            document_id=document_id,
            content_type=ContentType.TABLE,
            text=md_table,
            metadata={"rows": len(rows), "columns": len(rows[0]) if rows else 0}
        )

    def _table_to_markdown(self, rows: List[List[str]]) -> str:
        """Convert table rows to markdown format."""
        if not rows:
            return ""

        lines = []
        header = rows[0]

        # Header row
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # Data rows
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[:len(header)]) + " |")

        return "\n".join(lines)

    # =====================
    # DOC Extraction
    # =====================

    def _extract_doc(
        self,
        file_path: Path,
        document_id: str,
        include_tables: bool
    ) -> Tuple[List[ExtractedBlock], int]:
        """Extract text from DOC file using antiword or catdoc."""
        import subprocess

        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                text = result.stdout
                return self._parse_plain_text(text, document_id), None
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        try:
            result = subprocess.run(
                ['catdoc', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                text = result.stdout
                return self._parse_plain_text(text, document_id), None
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        raise RuntimeError(
            "Cannot extract DOC file. Install antiword or catdoc: "
            "sudo apt-get install antiword catdoc"
        )

    # =====================
    # TXT Extraction
    # =====================

    def _extract_txt(
        self,
        file_path: Path,
        document_id: str
    ) -> Tuple[List[ExtractedBlock], int]:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        return self._parse_plain_text(text, document_id), None

    def _parse_plain_text(self, text: str, document_id: str) -> List[ExtractedBlock]:
        """Parse plain text into structured blocks."""
        blocks = []
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            lines = para.split('\n')
            first_line = lines[0].strip()

            # Heading detection
            is_heading = (
                len(first_line) < 100 and
                (first_line.isupper() or
                 first_line.endswith(':') or
                 (len(lines) == 1 and not first_line.endswith('.')))
            )

            if is_heading and len(lines) == 1:
                blocks.append(self._create_block(
                    document_id=document_id,
                    content_type=ContentType.HEADING,
                    text=first_line,
                    level=2
                ))
            else:
                blocks.append(self._create_block(
                    document_id=document_id,
                    content_type=ContentType.PARAGRAPH,
                    text=para
                ))

        return blocks

    # =====================
    # Markdown Output
    # =====================

    def _build_markdown(self, blocks: List[ExtractedBlock]) -> str:
        """Build markdown formatted output from blocks."""
        lines = []
        current_page = None

        for block in blocks:
            # Page separator for PDFs
            if block.page and block.page != current_page:
                if current_page is not None:
                    lines.append("")
                    lines.append("---")
                    lines.append("")
                lines.append(f"## Page {block.page}")
                lines.append("")
                current_page = block.page

            # Format based on content type
            if block.content_type == ContentType.HEADING:
                prefix = "#" * (block.level or 2)
                lines.append(f"{prefix} {block.text}")
                lines.append("")

            elif block.content_type == ContentType.LIST_ITEM:
                lines.append(f"- {block.text}")

            elif block.content_type == ContentType.TABLE:
                lines.append("")
                lines.append(block.text)
                lines.append("")

            elif block.content_type == ContentType.IMAGE:
                lines.append("")
                lines.append(block.text)
                lines.append("")

            else:  # PARAGRAPH or TEXT
                lines.append(block.text)
                lines.append("")

        return "\n".join(lines).strip()
